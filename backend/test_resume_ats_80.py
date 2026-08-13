import os
import sys
from datetime import datetime
from docx import Document

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from app.ai.resume_parser import resume_parser
from app.ai.ats_scorer import ats_scorer

def create_sample_docx(file_path: str):
    """Programmatically generate a high-quality resume docx containing targeted developer keywords."""
    doc = Document()
    doc.add_heading("RAHUL SHARMA - AI ENGINEER & FULL-STACK DEVELOPER", level=1)
    
    doc.add_paragraph("Email: rahul.sharma@example.com | Phone: +919876543210 | Location: New Delhi, India")
    
    doc.add_heading("Professional Summary", level=2)
    doc.add_paragraph(
        "Highly skilled AI Engineer and Full-Stack Developer with 3+ years of experience designing, training, and "
        "deploying Machine Learning and NLP models. Proficient in FastAPI, Python, PostgreSQL, and Docker containerization."
    )
    
    doc.add_heading("Core Technical Skills", level=2)
    doc.add_paragraph(
        "Programming Languages: Python, Java, SQL, JavaScript\n"
        "AI/ML & NLP: Machine Learning, Deep Learning, Natural Language Processing, PyTorch, TensorFlow, scikit-learn, spaCy, Sentence Transformers\n"
        "Backend & Database: FastAPI, Flask, REST API, PostgreSQL, DBMS\n"
        "DevOps & Tools: Docker, AWS, Git, GitHub, Jenkins, Kubernetes"
    )
    
    doc.add_heading("Professional Experience", level=2)
    p1 = doc.add_paragraph("Senior AI Developer | TechCorp AI Solutions (2024 - Present)")
    doc.add_paragraph("- Built scalable REST APIs using FastAPI and PostgreSQL backend services.\n- Scaled model deployment pipelines using Docker containerization and AWS ECS.")
    
    p2 = doc.add_paragraph("Machine Learning Engineer | Innovate Lab (2022 - 2024)")
    doc.add_paragraph("- Fine-tuned BERT and Sentence Transformers models for semantic candidate search.\n- Optimized data ingestion with SQL queries.")
    
    doc.add_heading("Personal Projects", level=2)
    doc.add_paragraph("Project A: AI Unified Recruitment Platform (FastAPI, spaCy, Docker)\nProject B: LeetCode-style sandboxed execution environment\nProject C: Time-series skill demand forecasting pipeline")
    
    doc.add_heading("Education", level=2)
    doc.add_paragraph("Master of Technology in Computer Science (M.Tech) | IIT Delhi, 2022\nBachelor of Technology (B.Tech) | Delhi Technological University, 2020")
    
    doc.add_heading("Certifications", level=2)
    doc.add_paragraph("AWS Certified Solutions Architect | AWS\nCertified TensorFlow Developer | Google")

    doc.add_heading("Languages", level=2)
    doc.add_paragraph("English, Hindi")

    doc.save(file_path)
    print(f"[SUCCESS] Sample high-matching resume DOCX created at: {file_path}")

def run_ats_check():
    print("==================================================")
    print("RESUME ATS 80%+ COMPATIBILITY VERIFICATION")
    print("==================================================")

    docx_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "sample_ats_80_resume.docx")
    create_sample_docx(docx_path)

    # 1. Read the text from the generated DOCX file
    doc = Document(docx_path)
    lines = []
    for para in doc.paragraphs:
        if para.text.strip():
            lines.append(para.text)
    raw_text = "\n".join(lines)

    # 2. Parse the resume raw text using our parser
    parsed = resume_parser.parse(raw_text)

    # Define a target Job Description requiring AI/ML & backend skills
    job_skills = ["Python", "FastAPI", "Docker", "PostgreSQL", "PyTorch", "Git"]
    job_description = (
        "Looking for a Senior AI Engineer. Must be proficient in Python development, "
        "building REST APIs with FastAPI, managing databases on PostgreSQL, and deploying services via Docker. "
        "Experience in Deep Learning and PyTorch is essential."
    )

    # 3. Score the parsed resume against the job description
    ats_result = ats_scorer.score(parsed, job_description=job_description, job_skills=job_skills)

    print("\nParsed Resume Attributes:")
    print(f"  Name:     {parsed.get('name')}")
    print(f"  Email:    {parsed.get('email')}")
    print(f"  Skills:   {parsed.get('skills')}")
    print(f"  Projects: {len(parsed.get('projects'))} found")
    print(f"  Experience: {len(parsed.get('experience'))} found")

    print("\nATS Scoring Matrix Result:")
    print(f"  Overall Score: {ats_result['ats_score']}%")
    print(f"  Level Badge:   {ats_result['level']}")
    print(f"  Matched Skills: {ats_result['matched_skills']}")
    print(f"  Missing Skills: {ats_result['missing_skills']}")
    print(f"  Recommended Improvements: {ats_result.get('threshold_warning', {}).get('recommended_improvements', [])}")

    print("\nScore Breakdown:")
    for key, val in ats_result["score_breakdown"].items():
        print(f"  - {key:<20}: {val}%")

    print("\n==================================================")
    # Assert overall score is >= 80%
    if ats_result['ats_score'] >= 80.0:
        print("VERIFICATION SUCCESS: ATS Score is >= 80%!")
    else:
        print("VERIFICATION FAILURE: ATS Score is below 80%. Check weights or skills extraction.")
    print("==================================================")

if __name__ == "__main__":
    run_ats_check()
