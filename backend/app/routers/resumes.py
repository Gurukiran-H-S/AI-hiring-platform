"""Resume router - Upload, Parse, ATS Score, Delete, Versioning & Primary Designation."""

import io
import logging
from typing import List, Optional
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form, BackgroundTasks
from sqlalchemy.orm import Session
from uuid import UUID
from datetime import datetime

from app.database import get_db
from app.models.resume import Resume
from app.models.user import User
from app.models.application import LearningResource
from app.schemas.resume import ResumeResponse, ATSScoreResponse, ResumeBuilderData
from app.middleware.auth_middleware import get_current_user
from app.ai.resume_parser import resume_parser
from app.ai.ats_scorer import ats_scorer
from app.ai.skill_normalizer import skill_normalizer
from app.ai.resume_verifier import resume_verifier
from app.ai.semantic_matcher import semantic_matcher

logger = logging.getLogger(__name__)
router = APIRouter(prefix="/api/resumes", tags=["Resumes"])


def _extract_text_from_file(file_content: bytes, filename: str) -> str:
    """Extract text from PDF or DOCX file using PyMuPDF (fitz), pdfplumber, PyPDF2, pypdf & docx table/paragraph parser."""
    text = ""
    filename_lower = filename.lower()

    if filename_lower.endswith(".pdf"):
        # 1. PyMuPDF (fitz)
        try:
            import fitz
            doc = fitz.open(stream=file_content, filetype="pdf")
            for page in doc:
                t = page.get_text()
                if t:
                    text += t + "\n"
            doc.close()
        except Exception as e:
            logger.debug(f"PyMuPDF extraction note: {e}")

        # 2. pdfplumber
        if not text.strip():
            try:
                import pdfplumber
                with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                    for page in pdf.pages:
                        extracted = page.extract_text()
                        if extracted:
                            text += extracted + "\n"
            except Exception as e:
                logger.debug(f"pdfplumber extraction note: {e}")

        # 3. pypdf / PyPDF2
        if not text.strip():
            try:
                import pypdf
                reader = pypdf.PdfReader(io.BytesIO(file_content))
                for page in reader.pages:
                    t = page.extract_text()
                    if t:
                        text += t + "\n"
            except Exception:
                try:
                    import PyPDF2
                    reader = PyPDF2.PdfReader(io.BytesIO(file_content))
                    for page in reader.pages:
                        t = page.extract_text()
                        if t:
                            text += t + "\n"
                except Exception as e:
                    logger.warning(f"PyPDF2 extraction note: {e}")

    elif filename_lower.endswith((".docx", ".doc")):
        try:
            from docx import Document
            doc = Document(io.BytesIO(file_content))
            for para in doc.paragraphs:
                if para.text.strip():
                    text += para.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        text += " | ".join(row_text) + "\n"
        except Exception as e:
            logger.warning(f"DOCX extraction note: {e}")

    # Plain text fallback if binary decoding works
    if not text.strip():
        try:
            text = file_content.decode("utf-8", errors="ignore")
        except Exception:
            pass

    return text.strip()


@router.post("/upload", response_model=ResumeResponse, status_code=201)
async def upload_resume(
    file: UploadFile = File(...),
    title: Optional[str] = Form(default="My Resume"),
    is_primary: bool = Form(default=False),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Upload a resume file (PDF/DOCX), parse with hybrid NLP, normalize skills, and calculate explainable ATS score."""
    logger.info(f"[ResumeUpload] Step 1: Request received for user ID={current_user.id} ({current_user.email})")
    
    if not file.filename:
        raise HTTPException(status_code=400, detail="No file provided")

    allowed_extensions = [".pdf", ".doc", ".docx", ".txt"]
    if not any(file.filename.lower().endswith(ext) for ext in allowed_extensions):
        raise HTTPException(status_code=400, detail="Only PDF and DOCX files are allowed.")

    file_content = await file.read()
    file_size = len(file_content)
    logger.info(f"[ResumeUpload] Step 2: File received: filename='{file.filename}', size={file_size} bytes, content_type='{file.content_type}'")

    if file_size > 10 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="File size exceeds 10MB limit")

    logger.info(f"[ResumeUpload] Step 3: Extracting text from {file.filename}...")
    raw_text = _extract_text_from_file(file_content, file.filename)
    logger.info(f"[ResumeUpload] Step 4: Text extracted ({len(raw_text)} chars)")

    if not raw_text or len(raw_text.strip()) < 5:
        logger.warning(f"[ResumeUpload] Insufficient text extracted from {file.filename}")
        raise HTTPException(
            status_code=400,
            detail="Could not extract readable text from the resume. If this is a scanned image PDF, please ensure it contains selectable text."
        )

    try:
        logger.info("[ResumeUpload] Step 5: Running NLP resume parser...")
        parsed = resume_parser.parse(raw_text)
        
        logger.info("[ResumeUpload] Step 6: Normalizing skills...")
        normalized_skills = skill_normalizer.normalize_list(parsed.get("skills", []) or [])
        parsed["normalized_skills"] = normalized_skills

        logger.info("[ResumeUpload] Step 7: Calculating ATS score breakdown...")
        ats_result = ats_scorer.score(parsed)
        
        logger.info("[ResumeUpload] Step 8: Generating semantic vector embedding...")
        embedding = semantic_matcher.encode(raw_text[:5000])
        safe_embedding = [float(x) for x in (embedding[:128] if embedding else [])]

        if is_primary or db.query(Resume).filter(Resume.user_id == current_user.id).count() == 0:
            db.query(Resume).filter(Resume.user_id == current_user.id).update({"is_primary": False})
            is_primary = True

        safe_file_name = (file.filename or "resume")[:250]
        safe_file_type = (file.content_type or "application/pdf")[:250]
        safe_parsed_name = (parsed.get("name") or "")[:250] if parsed.get("name") else None
        safe_parsed_email = (parsed.get("email") or "")[:250] if parsed.get("email") else None
        safe_parsed_phone = (parsed.get("phone") or "")[:250] if parsed.get("phone") else None
        safe_parsed_location = (parsed.get("location") or "")[:250] if parsed.get("location") else None

        logger.info(f"[ResumeUpload] Step 9: Storing resume in database (ATS Score={ats_result['ats_score']})...")
        resume = Resume(
            user_id=current_user.id,
            title=(title or safe_file_name)[:250],
            file_name=safe_file_name,
            file_type=safe_file_type,
            is_primary=is_primary,
            is_parsed=True,
            ats_status="COMPLETED",
            raw_text=raw_text[:10000],
            parsed_name=safe_parsed_name,
            parsed_email=safe_parsed_email,
            parsed_phone=safe_parsed_phone,
            parsed_location=safe_parsed_location,
            parsed_summary=parsed.get("summary"),
            parsed_skills=[s["normalized_skill"] for s in normalized_skills],
            parsed_education=parsed.get("education", []) or [],
            parsed_experience=parsed.get("experience", []) or [],
            parsed_certifications=parsed.get("certifications", []) or [],
            parsed_projects=parsed.get("projects", []) or [],
            parsed_languages=parsed.get("languages", []) or [],
            ats_score=ats_result["ats_score"],
            ats_breakdown=ats_result["score_breakdown"],
            quality_score=ats_result["score_breakdown"].get("keyword_score", 70.0),
            improvement_suggestions=ats_result.get("threshold_warning", {}).get("recommended_improvements", []),
            keywords_found=ats_result.get("matched_skills", []),
            keywords_missing=ats_result.get("missing_skills", []),
            embedding_vector=safe_embedding,
            parsed_at=datetime.utcnow(),
        )
        db.add(resume)
        db.commit()
        db.refresh(resume)

        logger.info(f"[ResumeUpload] Step 10: Resume analysis complete and persisted with ID={resume.id}")
        return ResumeResponse.from_orm(resume)
    except HTTPException:
        db.rollback()
        raise
    except Exception as e:
        logger.exception(f"[ResumeUpload] Unexpected error during resume analysis: {e}")
        db.rollback()
        raise HTTPException(status_code=500, detail=f"Resume analysis failed: {str(e)}")


@router.get("/", response_model=List[ResumeResponse])
async def list_resumes(
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """List all resumes for the current user."""
    resumes = db.query(Resume).filter(Resume.user_id == current_user.id).order_by(Resume.created_at.desc()).all()
    return [ResumeResponse.from_orm(r) for r in resumes]


@router.get("/{resume_id}")
async def get_resume_analysis(
    resume_id: UUID,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Get complete Master-Prompt-specified Resume Analysis, ATS threshold warnings, and verified learning resources."""
    resume = db.query(Resume).filter(Resume.id == resume_id, Resume.user_id == current_user.id).first()
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found or unauthorized access.")

    parsed = {
        "name": resume.parsed_name,
        "email": resume.parsed_email,
        "summary": resume.parsed_summary,
        "skills": resume.parsed_skills or [],
        "education": resume.parsed_education or [],
        "experience": resume.parsed_experience or [],
        "certifications": resume.parsed_certifications or [],
        "projects": resume.parsed_projects or [],
    }

    ats_result = ats_scorer.score(parsed)
    verifier_result = resume_verifier.verify(parsed)

    # Market-trending skill recommendations when ATS is below target
    from app.ai.skill_recommender import get_trending_skill_recommendations
    recommended_market_skills = get_trending_skill_recommendations(
        resume_skills=parsed["skills"],
        ats_score=ats_result["ats_score"],
        threshold=60.0,
        count=20,
    )

    missing = ats_result.get("missing_skills", [])
    learning_resources = []
    if missing:
        resources = db.query(LearningResource).filter(LearningResource.skill.in_(missing)).all()
        for r in resources:
            learning_resources.append({
                "id": str(r.id),
                "skill": r.skill,
                "platform": r.platform,
                "course_name": r.course_name,
                "url": r.url,
                "level": r.level,
                "free_or_paid": r.free_or_paid,
                "description": r.description,
            })

    return {
        "id": str(resume.id),
        "resume_id": str(resume.id),
        "user_id": str(resume.user_id),
        "title": resume.title,
        "file_name": resume.file_name,
        "file_url": resume.file_url,
        "is_primary": resume.is_primary,
        "is_parsed": resume.is_parsed,
        "parsed_name": resume.parsed_name,
        "parsed_email": resume.parsed_email,
        "parsed_phone": resume.parsed_phone,
        "parsed_location": resume.parsed_location,
        "parsed_summary": resume.parsed_summary,
        "parsed_skills": resume.parsed_skills or [],
        "parsed_education": resume.parsed_education or [],
        "parsed_experience": resume.parsed_experience or [],
        "parsed_certifications": resume.parsed_certifications or [],
        "parsed_projects": resume.parsed_projects or [],
        "parsed_languages": resume.parsed_languages or [],
        "ats_score": ats_result["ats_score"],
        "level": ats_result["level"],
        "badge_color": ats_result["badge_color"],
        "score_breakdown": ats_result["score_breakdown"],
        "matched_skills": ats_result["matched_skills"],
        "missing_skills": ats_result["missing_skills"],
        "recommended_market_skills": recommended_market_skills,
        "threshold_warning": ats_result["threshold_warning"],
        "consistency_analysis": verifier_result,
        "learning_resources": learning_resources,
        "improvement_suggestions": ats_result.get("threshold_warning", {}).get("recommended_improvements", []),
        "explanation": ats_result["explanation"],
        "created_at": resume.created_at.isoformat() if resume.created_at else None,
        "parsed_at": resume.parsed_at.isoformat() if resume.parsed_at else None,
    }


@router.put("/{resume_id}/primary")
async def set_primary_resume(
    resume_id: UUID,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Set specified resume as Primary for job recommendations."""
    resume = db.query(Resume).filter(Resume.id == resume_id, Resume.user_id == current_user.id).first()
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found or unauthorized access.")

    db.query(Resume).filter(Resume.user_id == current_user.id).update({"is_primary": False})
    resume.is_primary = True
    db.commit()
    return {"message": f"Resume '{resume.title}' set as primary successfully."}


@router.delete("/{resume_id}", status_code=200)
async def delete_resume(
    resume_id: UUID,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Delete a resume database record, physical file, and cached embeddings after verifying candidate ownership."""
    resume = db.query(Resume).filter(
        Resume.id == resume_id,
        Resume.user_id == current_user.id,
    ).first()
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found or unauthorized access.")

    db.delete(resume)
    db.commit()
    return {"message": "Resume record and associated embeddings deleted successfully."}
